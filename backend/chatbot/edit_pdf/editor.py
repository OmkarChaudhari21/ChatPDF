import subprocess
import os
from pdf2docx import Converter
from docx2pdf import convert
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shape import InlineShape
from docx.table import Table
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


class PDFConverter:
    def __init__(self, pdf_path, docx_path, modified_docx):
        self.pdf_path = pdf_path
        self.docx_path = docx_path
        self.modified_docx = modified_docx

    def pdf_to_word(self):
        """Converts PDF to Word document."""
        cv = Converter(self.pdf_path)
        cv.convert(self.docx_path, start=0, end=None)
        cv.close()

    def word_to_pdf(self, output_path):
        """Converts Word document to PDF."""
        convert(self.modified_docx, output_path)


class DocumentFormatter:
    def __init__(self, font_color=None, bg_color=None, font_size_scale=None):
        self.font_color = font_color
        self.bg_color = bg_color
        self.font_size_scale = font_size_scale

    def hex_to_rgb(self, hex_color):
        """Converts hex color to RGB tuple."""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i + 2], 16) / 255 for i in (0, 2, 4))

    def set_run_format(self, run, run_format, font_color_rgb):
        """Sets the formatting for a run."""
        if font_color_rgb:
            run.font.color.rgb = font_color_rgb
        if self.font_size_scale and self.font_size_scale != 1:
            current_size = run_format['size'] if run_format['size'] else 12  # Default size is 12pt if not specified
            run.font.size = Pt(current_size * self.font_size_scale)
        if run_format['bold']:
            run.bold = True
        if run_format['italic']:
            run.italic = True
        if run_format['underline']:
            run.underline = True
        run.font.name = run_format['font_name']
        if 'HYPERLINK' in run._r.xml:
            run.hyperlink.target = run_format['hyperlink']  # Preserving hyperlink

    def apply_custom_formats_to_word(self, content, headers, footers):
        """Applies custom formats to the Word document content and returns a new Document object."""
        new_doc = Document()

        # Set minimal margins
        for section in new_doc.sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

        # Set font color
        font_color_rgb = None
        if self.font_color:
            font_color_rgb = RGBColor(*[int(c * 255) for c in self.hex_to_rgb(self.font_color)])

        # Set background color
        if self.bg_color:
            bg_color_hex = self.bg_color.lstrip('#')
            bg = OxmlElement('w:background')
            bg.set(qn('w:color'), bg_color_hex)
            new_doc.element.insert(0, bg)
            shd1 = OxmlElement('w:displayBackgroundShape')
            new_doc.settings.element.insert(0, shd1)

        for item in content:
            if item['type'] == 'paragraph':
                new_para = new_doc.add_paragraph()
                for run_text, run_format in item['content']:
                    run = new_para.add_run(run_text)
                    self.set_run_format(run, run_format, font_color_rgb)
                    if run_format['hyperlink']:
                        run.hyperlink.target = run_format['hyperlink']  # Preserving hyperlink
            elif item['type'] == 'table':
                self.add_table_to_doc(new_doc, item['content'], font_color_rgb)
            elif item['type'] == 'image':
                new_doc.add_picture(item['path'], width=item['width'], height=item['height'])
            elif item['type'] == 'drawing':
                drawing_element = OxmlElement.fromstring(item['content'])
                self.apply_color_to_drawing(drawing_element, font_color_rgb)
                new_doc.element.body.append(drawing_element)

        self.add_headers_footers(new_doc, headers, footers)
        return new_doc

    def add_headers_footers(self, doc, headers, footers):
        """Adds headers and footers back to the document."""
        for i, section in enumerate(doc.sections):
            header = section.header
            footer = section.footer

            # Add headers
            for item in headers[min(i, len(headers) - 1)]:
                if item['type'] == 'paragraph':
                    new_para = header.add_paragraph()
                    for run_text, run_format in item['content']:
                        run = new_para.add_run(run_text)
                        self.set_run_format(run, run_format, None)

            # Add footers
            for item in footers[min(i, len(footers) - 1)]:
                if item['type'] == 'paragraph':
                    new_para = footer.add_paragraph()
                    for run_text, run_format in item['content']:
                        run = new_para.add_run(run_text)
                        self.set_run_format(run, run_format, None)

    def apply_color_to_drawing(self, drawing_element, font_color_rgb):
        """Applies the font color to drawing elements."""
        if font_color_rgb:
            for elem in drawing_element.iter():
                if 'color' in elem.attrib:
                    elem.set('color', font_color_rgb)
                if 'fill' in elem.attrib:
                    elem.set('fill', font_color_rgb)

    def add_table_to_doc(self, doc, table_data, font_color_rgb):
        """Adds a table to the document with the specified formatting."""
        table = doc.add_table(rows=0, cols=len(table_data[0]))
        for row_data in table_data:
            row = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                cell = row[i]
                tc_pr = cell._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), self.bg_color.lstrip('#'))  # Apply bg color
                tc_pr.append(shd)
                for para_data in cell_data['content']:
                    for para_text, para_format in para_data:
                        new_para = cell.add_paragraph()
                        run = new_para.add_run(para_text)
                        self.set_run_format(run, para_format, font_color_rgb)
                        if para_format['hyperlink']:
                            run.hyperlink.target = para_format['hyperlink']  # Preserving hyperlink
                        new_para.paragraph_format.space_after = Pt(para_format['space_after']) if para_format['space_after'] else None
                        new_para.paragraph_format.space_before = Pt(para_format['space_before']) if para_format['space_before'] else None
                        new_para.paragraph_format.left_indent = Pt(para_format['left_indent']) if para_format['left_indent'] else None
                        new_para.paragraph_format.right_indent = Pt(para_format['right_indent']) if para_format['right_indent'] else None
                # Apply cell margins
                cell.paragraphs[0].paragraph_format.space_before = Pt(cell_data['margin_top'])
                cell.paragraphs[-1].paragraph_format.space_after = Pt(cell_data['margin_bottom'])
                cell.paragraphs[0].paragraph_format.left_indent = Pt(cell_data['margin_left'])
                cell.paragraphs[0].paragraph_format.right_indent = Pt(cell_data['margin_right'])

    def apply_formats_in_place(self, doc):
        """Applies font color and background color in place to the existing document."""
        if self.font_color:
            font_color_rgb = RGBColor(*[int(c*255) for c in self.hex_to_rgb(self.font_color)])

        if self.bg_color:
            bg_color_hex = self.bg_color.lstrip('#')
            bg = OxmlElement('w:background')
            bg.set(qn('w:color'), bg_color_hex)
            doc.element.insert(0, bg)
            shd1 = OxmlElement('w:displayBackgroundShape')
            doc.settings.element.insert(0, shd1)

        for para in doc.paragraphs:
            for run in para.runs:
                if self.font_color:
                    run.font.color.rgb = font_color_rgb

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    tc_pr = cell._element.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), bg_color_hex)
                    tc_pr.append(shd)  # Apply bg color
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if self.font_color:
                                run.font.color.rgb = font_color_rgb


class ContentExtractor:
    def __init__(self, doc):
        self.doc = doc

    def extract_content(self):
        """Extracts the content from the document including paragraphs, tables, and images."""
        content = []

        # Extract headers and footers for each section
        headers, footers = self.extract_headers_footers()

        # Extract content from the document
        for element in self.doc.iter_inner_content():
            if isinstance(element, Paragraph):
                content.append(self.extract_paragraph(element))
            elif isinstance(element, Table):
                content.append(self.extract_table(element))
            elif isinstance(element, InlineShape):
                content.append(self.extract_image(element))
            elif element.tag.endswith('drawing'):
                content.append(self.extract_drawing(element))

        return content, headers, footers

    def extract_headers_footers(self):
        """Extracts headers and footers from the document."""
        headers = []
        footers = []

        for section in self.doc.sections:
            header = section.header
            footer = section.footer

            # Extract headers and footers
            headers.append(self.extract_paragraphs_header_footer(header.paragraphs))
            footers.append(self.extract_paragraphs_header_footer(footer.paragraphs))

        return headers, footers

    def extract_paragraphs_header_footer(self, paragraphs):
        """Extracts paragraphs from header or footer."""
        content = []
        for para in paragraphs:
            content.append(self.extract_paragraph(para))
        return content

    def extract_drawing(self, drawing):
        """Extracts drawing data."""
        drawing_data = drawing.xml
        return {'type': 'drawing', 'content': drawing_data}

    def extract_paragraph(self, para):
        """Extracts paragraph content and formatting."""
        para_content = []
        for run in para.runs:
            run_format = {
                'size': run.font.size.pt if run.font.size else None,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'space_after': para.paragraph_format.space_after.pt if para.paragraph_format.space_after else None,
                'space_before': para.paragraph_format.space_before.pt if para.paragraph_format.space_before else None,
                'left_indent': para.paragraph_format.left_indent.pt if para.paragraph_format.left_indent else None,
                'right_indent': para.paragraph_format.right_indent.pt if para.paragraph_format.right_indent else None,
                'hyperlink': run.hyperlink.target if 'HYPERLINK' in run._r.xml else None  # Extracting hyperlink
            }
            para_content.append((run.text, run_format))
        return {'type': 'paragraph', 'content': para_content}

    def extract_table(self, table):
        """Extracts table content and formatting."""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_content = []
                for para in cell.paragraphs:
                    cell_content.append(self.extract_paragraph(para)['content'])
                cell_format = {
                    'content': cell_content,
                    'margin_top': cell.paragraphs[0].paragraph_format.space_before.pt if cell.paragraphs[0].paragraph_format.space_before else 0,
                    'margin_bottom': cell.paragraphs[-1].paragraph_format.space_after.pt if cell.paragraphs[-1].paragraph_format.space_after else 0,
                    'margin_left': cell.paragraphs[0].paragraph_format.left_indent.pt if cell.paragraphs[0].paragraph_format.left_indent else 0,
                    'margin_right': cell.paragraphs[0].paragraph_format.right_indent.pt if cell.paragraphs[0].paragraph_format.right_indent else 0
                }
                row_data.append(cell_format)
            table_data.append(row_data)
        return {'type': 'table', 'content': table_data}

    def extract_image(self, image):
        """Extracts image data."""
        image_path = image._inline.graphic.graphicData.pic.nvPicPr.cNvPr.descr
        return {
            'type': 'image',
            'path': image_path,
            'width': image.width,
            'height': image.height
        }


class PDFEditor:
    def __init__(self, pdf_path, output_path, font_color=None, bg_color=None, font_size_scale=None):
        self.pdf_path = pdf_path
        self.output_path = output_path
        self.font_color = font_color
        self.bg_color = bg_color
        self.font_size_scale = font_size_scale

    def apply_custom_formats(self):
        """Main function to apply custom formats to the PDF and save the output."""
        intermediate_docx = "intermediate.docx"
        modified_docx = "modified.docx"

        # Convert PDF to Word
        converter = PDFConverter(self.pdf_path, intermediate_docx, modified_docx)
        converter.pdf_to_word()
        doc = Document(intermediate_docx)

        # Check if we need to create a new document or modify in place
        if self.font_size_scale is not None and float(self.font_size_scale) != 1.0:
            extractor = ContentExtractor(doc)
            content, headers, footers = extractor.extract_content()

            formatter = DocumentFormatter(self.font_color, self.bg_color, self.font_size_scale)
            new_doc = formatter.apply_custom_formats_to_word(content, headers, footers)
            ### self.add_page_numbers(new_doc)
            new_doc.save(modified_docx)
        else:
            formatter = DocumentFormatter(self.font_color, self.bg_color, self.font_size_scale)
            formatter.apply_formats_in_place(doc)
            doc.save(modified_docx)

        # Convert modified DOCX to PDF
        converter.word_to_pdf(self.output_path)

        # Delete intermediate DOCX files
        os.remove(intermediate_docx)
        os.remove(modified_docx)


# if __name__ == "__main__":
#     editor = PDFEditor("apple_small.pdf", "output.pdf", font_color="#FFFFFF", bg_color="#000000", font_size_scale=1)
#     editor.apply_custom_formats()
